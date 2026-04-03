import mimetypes
import os
import re
from pathlib import Path
from typing import Any, Optional
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader

from langchain_core.messages import HumanMessage, SystemMessage
from langchain_gigachat.chat_models import GigaChat

from app.core.config import settings
from app.core.prompts import get_document_prompt
from app.services.gigachat_client import get_gigachat_client
from app.utils.text import capitalize_filename


class DocumentGenerationService:
    def __init__(self):
        self.model = settings.GIGACHAT_MODEL or "GigaChat-Pro"
        self.temperature = 0.1

    async def _get_llm(self) -> GigaChat:
        client = await get_gigachat_client()
        token = await client.get_valid_token()
        return GigaChat(
            access_token=token,
            model=self.model,
            temperature=self.temperature,
            verify_ssl_certs=False,  # Для разработки
        )

    def _safe_filename(self, name: str) -> str:
        name = name.strip()
        name = re.sub(r"[\\/:*?\"<>|]+", "_", name)
        return name or f"document_{uuid4().hex}"

    def _truncate(self, text: str, max_chars: int) -> str:
        text = text or ""
        if len(text) <= max_chars:
            return text
        return text[:max_chars] + "\n…(truncated)…"

    def extract_template_content(self, template_path: str, mime_type: Optional[str] = None) -> str:
        """
        Извлекает текст из шаблона для прокидывания в промпт.
        Для PDF используем pypdf (может плохо извлекать текст из сканов).
        """
        path = Path(template_path)
        suffix = path.suffix.lower()

        try:
            if suffix == ".pdf":
                print(f"[doc_gen] extract pdf template pages from {path.name}")
                reader = PdfReader(str(path))
                # Ограничиваем чтение PDF-шаблона по страницам, чтобы не зависать на больших/сканированных файлах.
                max_pages = 3
                max_chars = 12_000
                parts: list[str] = []
                for i, page in enumerate(reader.pages):
                    if i >= max_pages:
                        break
                    try:
                        page_text = page.extract_text() or ""
                    except Exception:
                        page_text = ""
                    if page_text:
                        parts.append(page_text)
                    current_len = sum(len(p) for p in parts)
                    if current_len >= max_chars:
                        break
                content = "\n".join(parts).strip()
                return self._truncate(content, max_chars)

            if suffix in (".docx",):
                print(f"[doc_gen] extract docx template from {path.name}")
                doc = DocxDocument(str(path))
                paragraphs = [p.text for p in doc.paragraphs if (p.text or "").strip()]
                # Чтобы не раздувать промпт, берём первые 60 абзацев.
                paragraphs = paragraphs[:60]
                content = "\n".join(paragraphs).strip()
                return self._truncate(content, 12_000)

            if suffix in (".txt",):
                print(f"[doc_gen] extract txt template from {path.name}")
                return path.read_text(encoding="utf-8", errors="ignore").strip()
        except Exception:
            # Если извлечение текста не удалось, возвращаем пустую строку,
            # чтобы модель всё равно смогла сгенерировать документ по описанию.
            return ""

        # Fallback: если неизвестное расширение
        return ""

    async def generate_document_text(
        self,
        *,
        user_query: str,
        document_type: str,
        template_name: str,
        template_content: str,
        client_data: dict[str, Any],
    ) -> str:
        system_prompt = get_document_prompt(
            document_type=document_type,
            template_name=template_name,
            user_query=user_query,
            template_content=self._truncate(template_content, 9000),
            client_data=client_data,
        )

        llm = await self._get_llm()
        response = await llm.ainvoke(
            [
                SystemMessage(content=system_prompt),
                HumanMessage(content=user_query),
            ]
        )
        return response.content or ""

    def render_docx_from_text(self, generated_text: str, output_docx_path: str) -> None:
        doc = DocxDocument()

        # Простая разметка: пустая строка -> новый абзац.
        normalized = (generated_text or "").replace("\r\n", "\n")
        paragraphs = re.split(r"\n\s*\n", normalized.strip()) if normalized.strip() else []
        if not paragraphs:
            doc.add_paragraph("")
        else:
            for block in paragraphs:
                block = block.strip("\n")
                if not block:
                    doc.add_paragraph("")
                    continue
                # Сохраняем переносы строк как отдельные строки внутри абзаца (простая аппроксимация).
                lines = block.split("\n")
                first = True
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    if first:
                        doc.add_paragraph(line)
                        first = False
                    else:
                        doc.add_paragraph(line)

        output_path = Path(output_docx_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

    def convert_docx_to_pdf(self, output_docx_path: str, output_pdf_path: str) -> None:
        """
        Конвертация DOCX->PDF.
        docx2pdf на Windows требует Word/совместимую подсистему.
        """
        # Импортируем динамически, чтобы сервис мог стартовать даже без docx2pdf
        from docx2pdf import convert  # type: ignore

        pdf_path = Path(output_pdf_path)
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        convert(output_docx_path, str(pdf_path))

    async def generate_document_files(
        self,
        *,
        user_query: str,
        document_type: str,
        template_name: str,
        template_path: Optional[str],
        output_format: str,
        client_data: dict[str, Any],
        output_dir: str,
        provided_template_content: Optional[str] = None,
    ) -> dict[str, Any]:
        template_content = provided_template_content or ""
        if not template_content and template_path:
            template_content = self.extract_template_content(template_path)

        generated_text = await self.generate_document_text(
            user_query=user_query,
            document_type=document_type,
            template_name=template_name,
            template_content=template_content,
            client_data=client_data,
        )
        print(f"[doc_gen] generated_text ready len={len(generated_text or '')}")

        output_dir_path = Path(output_dir)
        output_dir_path.mkdir(parents=True, exist_ok=True)

        base_name = self._safe_filename(f"{document_type}_{uuid4().hex[:8]}")
        output_docx_path = str(output_dir_path / f"{base_name}.docx")
        self.render_docx_from_text(generated_text, output_docx_path)

        if output_format == "pdf":
            output_pdf_path = str(output_dir_path / f"{base_name}.pdf")
            try:
                self.convert_docx_to_pdf(output_docx_path, output_pdf_path)
                file_path = output_pdf_path
                mime_type = "application/pdf"
            except Exception:
                # Если конвертация в PDF не удалась (Word/COM/другая ошибка) —
                # не ломаем UX, отдаём DOCX.
                file_path = output_docx_path
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            file_path = output_docx_path
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        file_size = Path(file_path).stat().st_size
        return {
            "generated_text": generated_text,
            "file_path": file_path,
            "file_size": file_size,
            "mime_type": mime_type,
            "title": capitalize_filename(os.path.basename(file_path)),
        }

