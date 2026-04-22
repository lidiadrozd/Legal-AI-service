from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import re
from string import Formatter
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


@dataclass(frozen=True)
class TemplateField:
    key: str
    label: str
    multiline: bool = False
    pattern: str | None = None


@dataclass(frozen=True)
class DocumentTemplate:
    key: str
    version: int
    title: str
    description: str
    body: str
    fields: tuple[TemplateField, ...]


@dataclass(frozen=True)
class RenderedTemplate:
    key: str
    title: str
    description: str
    rendered_text: str


TEMPLATES: dict[str, DocumentTemplate] = {
    "statement_of_claim_arbitration": DocumentTemplate(
        key="statement_of_claim_arbitration",
        version=1,
        title="Исковое заявление в арбитражный суд",
        description="Базовый шаблон иска о взыскании задолженности.",
        body=(
            "В {court_name}\n"
            "Истец: {plaintiff_name}, адрес: {plaintiff_address}\n"
            "Ответчик: {defendant_name}, адрес: {defendant_address}\n"
            "Цена иска: {claim_amount}\n\n"
            "ИСКОВОЕ ЗАЯВЛЕНИЕ\n"
            "о взыскании задолженности\n\n"
            "Обстоятельства дела:\n{facts}\n\n"
            "Правовое обоснование:\n{legal_basis}\n\n"
            "На основании изложенного прошу суд:\n{requests}\n\n"
            "Приложения:\n{attachments}\n\n"
            "Дата: {date}\n"
            "Подпись: {plaintiff_name}\n"
        ),
        fields=(
            TemplateField("court_name", "Суд"),
            TemplateField("plaintiff_name", "Истец"),
            TemplateField("plaintiff_address", "Адрес истца"),
            TemplateField("defendant_name", "Ответчик"),
            TemplateField("defendant_address", "Адрес ответчика"),
            TemplateField("claim_amount", "Цена иска", pattern=r".*\d.*"),
            TemplateField("facts", "Обстоятельства дела", multiline=True),
            TemplateField("legal_basis", "Правовое обоснование", multiline=True),
            TemplateField("requests", "Требования к суду", multiline=True),
            TemplateField("attachments", "Приложения", multiline=True),
            TemplateField("date", "Дата", pattern=r"^\d{4}-\d{2}-\d{2}$"),
        ),
    ),
    "motion_to_postpone_hearing": DocumentTemplate(
        key="motion_to_postpone_hearing",
        version=1,
        title="Ходатайство об отложении судебного заседания",
        description="Шаблон процессуального ходатайства.",
        body=(
            "В {court_name}\n"
            "По делу № {case_number}\n"
            "Заявитель: {applicant_name}\n\n"
            "ХОДАТАЙСТВО\n"
            "об отложении судебного заседания\n\n"
            "Уважаемый суд!\n"
            "Прошу отложить судебное заседание, назначенное на {hearing_date},\n"
            "по следующим причинам:\n{reasons}\n\n"
            "Предлагаемая дата после отложения: {new_hearing_date}\n\n"
            "Приложения:\n{attachments}\n\n"
            "Дата: {date}\n"
            "Подпись: {applicant_name}\n"
        ),
        fields=(
            TemplateField("court_name", "Суд"),
            TemplateField("case_number", "Номер дела", pattern=r"^[АA]\d{1,2}-\d+/\d{4}$"),
            TemplateField("applicant_name", "Заявитель"),
            TemplateField("hearing_date", "Текущая дата заседания", pattern=r"^\d{4}-\d{2}-\d{2}$"),
            TemplateField("reasons", "Причины отложения", multiline=True),
            TemplateField("new_hearing_date", "Предлагаемая новая дата", pattern=r"^\d{4}-\d{2}-\d{2}$"),
            TemplateField("attachments", "Приложения", multiline=True),
            TemplateField("date", "Дата", pattern=r"^\d{4}-\d{2}-\d{2}$"),
        ),
    ),
    "appeal_complaint": DocumentTemplate(
        key="appeal_complaint",
        version=1,
        title="Апелляционная жалоба",
        description="Шаблон апелляционной жалобы по гражданскому/арбитражному делу.",
        body=(
            "В {appeal_court_name}\n"
            "Через {first_instance_court_name}\n"
            "Дело № {case_number}\n"
            "Заявитель: {appellant_name}\n\n"
            "АПЕЛЛЯЦИОННАЯ ЖАЛОБА\n\n"
            "Обжалуемый судебный акт: {challenged_act}\n\n"
            "Считаю решение незаконным и необоснованным по следующим основаниям:\n"
            "{grounds}\n\n"
            "Прошу суд апелляционной инстанции:\n{requests}\n\n"
            "Приложения:\n{attachments}\n\n"
            "Дата: {date}\n"
            "Подпись: {appellant_name}\n"
        ),
        fields=(
            TemplateField("appeal_court_name", "Суд апелляционной инстанции"),
            TemplateField("first_instance_court_name", "Суд первой инстанции"),
            TemplateField("case_number", "Номер дела", pattern=r"^[АA]\d{1,2}-\d+/\d{4}$"),
            TemplateField("appellant_name", "Заявитель"),
            TemplateField("challenged_act", "Обжалуемый акт"),
            TemplateField("grounds", "Основания жалобы", multiline=True),
            TemplateField("requests", "Требования", multiline=True),
            TemplateField("attachments", "Приложения", multiline=True),
            TemplateField("date", "Дата", pattern=r"^\d{4}-\d{2}-\d{2}$"),
        ),
    ),
}


def _required_fields(template_body: str) -> set[str]:
    fields: set[str] = set()
    for _, field_name, _, _ in Formatter().parse(template_body):
        if field_name:
            fields.add(field_name)
    return fields


def list_templates_meta() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for template in sorted(TEMPLATES.values(), key=lambda t: t.title):
        rows.append(
            {
                "key": template.key,
                "version": template.version,
                "title": template.title,
                "description": template.description,
                "fields": [
                    {
                        "key": field.key,
                        "label": field.label,
                        "multiline": field.multiline,
                        "pattern": field.pattern,
                    }
                    for field in template.fields
                ],
            }
        )
    return rows


def render_template(
    template_key: str,
    values: dict[str, Any],
    *,
    template_version: int | None = None,
) -> RenderedTemplate:
    template = TEMPLATES.get(template_key)
    if template is None:
        raise ValueError(f"Unknown template_key: {template_key}")
    if template_version is not None and template_version != template.version:
        raise ValueError(
            f"Template version mismatch: expected {template.version}, got {template_version}"
        )

    required = _required_fields(template.body)
    missing = sorted([field for field in required if str(values.get(field, "")).strip() == ""])
    if missing:
        raise KeyError(f"Missing required fields: {', '.join(missing)}")

    normalized_values = {k: str(v) for k, v in values.items()}
    for field in template.fields:
        if field.pattern is None:
            continue
        value = normalized_values.get(field.key, "").strip()
        if not re.fullmatch(field.pattern, value):
            raise ValueError(f"Field '{field.key}' has invalid format")

    rendered = template.body.format(**normalized_values)
    return RenderedTemplate(
        key=template.key,
        title=template.title,
        description=template.description,
        rendered_text=rendered,
    )


def build_docx_bytes(rendered: RenderedTemplate) -> bytes:
    doc = DocxDocument()
    title = doc.add_paragraph(rendered.title)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if title.runs:
        title.runs[0].bold = True

    for line in rendered.rendered_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        paragraph = doc.add_paragraph(stripped)
        if stripped.endswith(":") or stripped.isupper():
            for run in paragraph.runs:
                run.bold = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _resolve_pdf_font() -> tuple[str, bool]:
    candidate_fonts = [
        ("Arial", "C:/Windows/Fonts/arial.ttf"),
        ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for font_name, path in candidate_fonts:
        try:
            pdfmetrics.registerFont(TTFont(font_name, path))
            return font_name, True
        except Exception:
            continue
    return "Helvetica", False


def build_pdf_bytes(rendered: RenderedTemplate) -> bytes:
    font_name, has_unicode_font = _resolve_pdf_font()
    buf = BytesIO()
    pdf = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x_margin = 48
    y = height - 56

    pdf.setFont(font_name, 14)
    title = rendered.title if has_unicode_font else "Generated Legal Document"
    pdf.drawCentredString(width / 2, y, title)
    y -= 28

    pdf.setFont(font_name, 11)
    for original_line in rendered.rendered_text.splitlines():
        line = original_line.strip() or " "
        if not has_unicode_font:
            line = line.encode("latin-1", errors="replace").decode("latin-1")
        if y < 50:
            pdf.showPage()
            pdf.setFont(font_name, 11)
            y = height - 56
        pdf.drawString(x_margin, y, line[:130])
        y -= 16

    pdf.save()
    return buf.getvalue()
