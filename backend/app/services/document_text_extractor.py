"""
Извлечение текста из загруженных документов и фото для контекста чата.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from app.core.config import settings
from app.services.document_ocr import IMAGE_SUFFIXES, is_image_file, is_tesseract_available, ocr_image, ocr_pdf_pages

MAX_DOC_CHARS = 25_000
MAX_TOTAL_CHARS = 40_000
MIN_PDF_TEXT_CHARS = 20


def truncate_text(text: str, max_chars: int) -> str:
    cleaned = (text or "").strip()
    if len(cleaned) <= max_chars:
        return cleaned
    return cleaned[:max_chars] + "\n...[текст обрезан]"


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            parts.append(text)

    text_layer = "\n".join(parts).strip()
    if len(text_layer) >= MIN_PDF_TEXT_CHARS:
        return text_layer

    if not is_tesseract_available():
        if text_layer:
            return text_layer
        raise ValueError(
            "PDF без текстового слоя. На сервере не установлен OCR — загрузите фото или PDF со сканом после настройки Tesseract."
        )

    ocr_text = ocr_pdf_pages(
        path,
        lang=settings.OCR_LANGS,
        max_pages=settings.OCR_MAX_PDF_PAGES,
    )
    if ocr_text.strip():
        return ocr_text
    return text_layer


def _extract_image(path: Path) -> str:
    if not is_tesseract_available():
        raise ValueError(
            "OCR недоступен на сервере. Загрузите PDF/DOCX/TXT с текстовым слоем или обратитесь к администратору."
        )
    text = ocr_image(path, lang=settings.OCR_LANGS)
    if not text.strip():
        raise ValueError("На изображении не удалось распознать текст")
    return text


def extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".doc":
        raise ValueError("Формат .doc не поддерживается — сохраните файл как .docx")
    if is_image_file(path) or suffix in IMAGE_SUFFIXES:
        return _extract_image(path)
    raise ValueError(f"Неподдерживаемый формат файла: {suffix or '(без расширения)'}")


__all__ = [
    "MAX_DOC_CHARS",
    "MAX_TOTAL_CHARS",
    "extract_text_from_file",
    "truncate_text",
]
