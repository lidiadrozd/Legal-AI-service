"""
Заполнение загруженных .docx шаблонов с плейсхолдерами вида {{field_name}}.
Сохраняет структуру документа; при замене в абзаце текст пишется в первый run (ограничение python-docx при склейке runs).
"""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def _paragraph_full_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def _replace_paragraph_placeholders(paragraph: Paragraph, mapping: dict[str, str]) -> None:
    full = _paragraph_full_text(paragraph)
    if "{{" not in full:
        return
    new_full = full
    for key, value in mapping.items():
        new_full = re.sub(
            r"\{\{\s*" + re.escape(key) + r"\s*\}\}",
            value,
            new_full,
        )
    if new_full == full:
        return
    runs = list(paragraph.runs)
    if runs:
        runs[0].text = new_full
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(new_full)


def _walk_paragraphs_in_table(table: Table, fn) -> None:
    for row in table.rows:
        for cell in row.cells:
            _walk_block(cell._tc, cell, fn)


def _walk_block(container_elm, parent, fn) -> None:
    for child in container_elm.iterchildren():
        if child.tag == qn("w:p"):
            fn(Paragraph(child, parent))
        elif child.tag == qn("w:tbl"):
            _walk_paragraphs_in_table(Table(child, parent), fn)


def _walk_document_body(doc: DocumentObject, fn) -> None:
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            fn(Paragraph(child, doc))
        elif child.tag == qn("w:tbl"):
            _walk_paragraphs_in_table(Table(child, doc), fn)


def _walk_header_footer_part(part, fn) -> None:
    if part is None:
        return
    try:
        root = part._element
    except Exception:
        return
    for child in root.iterchildren():
        if child.tag == qn("w:p"):
            fn(Paragraph(child, part))
        elif child.tag == qn("w:tbl"):
            _walk_paragraphs_in_table(Table(child, part), fn)


def _iter_all_paragraphs(doc: DocumentObject) -> list[Paragraph]:
    collected: list[Paragraph] = []

    def collect(p: Paragraph) -> None:
        collected.append(p)

    _walk_document_body(doc, collect)
    for section in doc.sections:
        for attr in (
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ):
            try:
                hf = getattr(section, attr, None)
            except Exception:
                hf = None
            _walk_header_footer_part(hf, collect)
    return collected


def collect_placeholder_keys_from_docx(path: Path) -> list[str]:
    doc = Document(str(path))
    keys: set[str] = set()
    for p in _iter_all_paragraphs(doc):
        text = _paragraph_full_text(p)
        for m in _PLACEHOLDER_RE.finditer(text):
            keys.add(m.group(1))
    return sorted(keys)


def fill_docx_template(path: Path, fields: dict[str, Any]) -> bytes:
    """Читает шаблон с диска, подставляет значения, возвращает байты нового .docx."""
    mapping = {str(k): str(v) if v is not None else "" for k, v in fields.items()}
    doc = Document(str(path))
    for p in _iter_all_paragraphs(doc):
        _replace_paragraph_placeholders(p, mapping)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
